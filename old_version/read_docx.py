import docx
import os

files = [
    "FOR-PRO-18-9 Fiches produits types.docx",
    "FOR-PRO-Comp0500070-REV007.docx"
]

with open("output.txt", "w", encoding="utf-8") as f:
    for file in files:
        try:
            doc = docx.Document(file)
            f.write(f"=== File: {file} ===\n")
            
            f.write("--- Paragraphs ---\n")
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    f.write(text + "\n")
                    
            f.write("--- Tables ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip().replace("\n", " "))
                    if any(row_data):
                        f.write(" | ".join(row_data) + "\n")
            f.write("\n\n")
        except Exception as e:
            f.write(f"Error reading {file}: {e}\n")
